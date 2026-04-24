"""Generic raw transcript template."""

from docx import Document

try:
    from .template_common import add_item_image, add_para, add_table, style_document, title_from_source
except ImportError:
    from template_common import add_item_image, add_para, add_table, style_document, title_from_source


def render(extracted, plan, output):
    doc = Document()
    style_document(doc)
    source = extracted.get('source', '')
    doc.add_heading(title_from_source(source), level=0)
    add_para(doc, f'Source PPTX: {source}', size=9)
    add_para(doc, f"Conversion mode: generic_raw; fidelity: {plan.get('fidelity_level', 'F1')}", size=9)

    for slide in extracted.get('slides', []):
        doc.add_page_break()
        title = slide.get('title') or f"Slide {slide.get('slide_no')}"
        doc.add_heading(f"Slide {slide.get('slide_no')}: {title}", level=1)
        for item in slide.get('items', []):
            item_type = item.get('type')
            if item_type == 'text':
                style = 'List Bullet' if item.get('level', 0) > 0 else None
                doc.add_paragraph(item.get('text', ''), style=style)
            elif item_type == 'table':
                add_table(doc, item.get('rows', []))
            elif item_type == 'image':
                if not add_item_image(doc, item, plan.get('include_images')):
                    add_para(doc, '[Image recorded in source map; not embedded by selected plan.]', size=9)
            elif item.get('semantic_guess') == 'complex_visual':
                add_para(doc, '[Complex visual item recorded in fidelity ledger for manual review.]', size=9)
        if slide.get('notes'):
            doc.add_heading('Speaker Notes', level=2)
            add_para(doc, slide['notes'])

    doc.save(output)
    return output
