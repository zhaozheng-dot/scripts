"""Investment review DOCX template."""

import os
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLUE = '12355B'
TEAL = '1B7F79'
GOLD = 'D9A441'
LIGHT_BLUE = 'EAF4FB'
LIGHT_TEAL = 'E8F5F3'
LIGHT_GOLD = 'FFF4D8'
LIGHT_RED = 'FCE8E6'
LIGHT_GRAY = 'F4F6F8'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = cell._tc.get_or_add_tcPr().find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    for name in ['Normal', 'Title', 'Heading 1', 'Heading 2']:
        style = doc.styles[name]
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Title'].font.size = Pt(24)
    doc.styles['Title'].font.bold = True
    doc.styles['Title'].font.color.rgb = RGBColor.from_string(BLUE)
    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 1'].font.bold = True
    doc.styles['Heading 1'].font.color.rgb = RGBColor.from_string(BLUE)
    doc.styles['Heading 2'].font.size = Pt(12)
    doc.styles['Heading 2'].font.bold = True
    doc.styles['Heading 2'].font.color.rgb = RGBColor.from_string(TEAL)
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)


def para(doc, text='', bold=False, color=None, size=None, align=None, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Microsoft YaHei'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    return p


def add_heading(doc, no, title, subtitle=''):
    para(doc, f'{no}  {title}', bold=True, color=BLUE, size=16, after=4)
    if subtitle:
        para(doc, subtitle, color='64748B', size=9.5, after=8)


def add_badges(doc, badges):
    table = doc.add_table(rows=1, cols=len(badges))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, badge in zip(table.rows[0].cells, badges):
        set_cell_shading(cell, badge.get('fill', LIGHT_GRAY))
        set_cell_text(cell, badge['label'], bold=True, color=badge.get('color', BLUE), size=9.5)


def add_cards(doc, cards, cols=3):
    rows = (len(cards) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            if idx >= len(cards):
                continue
            card = cards[idx]
            set_cell_shading(cell, card.get('fill', LIGHT_GRAY))
            set_cell_text(cell, card['title'], bold=True, color=card.get('color', BLUE), size=9.5)
            p = cell.add_paragraph(card['body'])
            for run in p.runs:
                run.font.size = Pt(8.8)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            idx += 1


def add_table(doc, rows, header=None):
    if not rows:
        return
    offset = 1 if header else 0
    table = doc.add_table(rows=len(rows) + offset, cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        for i, value in enumerate(header):
            set_cell_text(table.rows[0].cells[i], value, bold=True, color='FFFFFF')
            set_cell_shading(table.rows[0].cells[i], BLUE)
    for r_idx, row in enumerate(rows, start=offset):
        for c_idx, value in enumerate(row):
            set_cell_text(table.rows[r_idx].cells[c_idx], value, bold=(c_idx == 0 and not header), size=9)
            if c_idx == 0 and not header:
                set_cell_shading(table.rows[r_idx].cells[c_idx], LIGHT_BLUE)


def add_callout(doc, title, body, fill=LIGHT_GOLD, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_text(cell, title, bold=True, color=color, size=10)
    p = cell.add_paragraph(body)
    for run in p.runs:
        run.font.size = Pt(9.2)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def collect_text(extracted):
    lines = []
    for slide in extracted.get('slides', []):
        for item in slide.get('items', []):
            if item.get('type') == 'text':
                lines.append(item.get('text', ''))
    return lines


def find_lines(lines, keys, limit=6):
    found = []
    for line in lines:
        if any(key.lower() in line.lower() for key in keys):
            if line not in found:
                found.append(line)
        if len(found) >= limit:
            break
    return found


def render(extracted, plan, output):
    doc = Document()
    style_doc(doc)
    source = extracted.get('source', '')
    title = os.path.splitext(os.path.basename(source))[0]
    lines = collect_text(extracted)

    para(doc, 'CONFIDENTIAL · INVESTMENT REVIEW', bold=True, color=TEAL, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    para(doc, title, bold=True, color=BLUE, size=24, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    para(doc, '专业报告版 | Office Agent 门控式转换', color='64748B', size=10, align=WD_ALIGN_PARAGRAPH.CENTER, after=16)
    add_badges(doc, [
        {'label': f"模式: {plan.get('selected_mode')}", 'fill': LIGHT_BLUE},
        {'label': f"忠实度: {plan.get('fidelity_level')}", 'fill': LIGHT_GOLD, 'color': '8A5A00'},
        {'label': f"风险级别: {plan.get('risk_level')}", 'fill': LIGHT_RED, 'color': 'B42318'},
    ])
    add_callout(doc, '转换边界', '本报告按确认方案进行轻度重组，不新增未经来源支持的事实。完整处理记录见 fidelity ledger。', fill='EEF2F7')
    doc.add_page_break()

    add_heading(doc, '01', '执行摘要')
    add_cards(doc, [
        {'title': '核心发现', 'body': '\n'.join(find_lines(lines, ['获批', '团队', '商业化'], 3)) or '见原始材料。', 'fill': LIGHT_TEAL, 'color': TEAL},
        {'title': '关键风险', 'body': '\n'.join(find_lines(lines, ['pending', '风险', '注册资本'], 3)) or '见原始材料。', 'fill': LIGHT_RED, 'color': 'B42318'},
        {'title': '建议动作', 'body': '\n'.join(find_lines(lines, ['建议', '尽调', '计划'], 3)) or '建议进入确认后的尽调流程。', 'fill': LIGHT_GOLD, 'color': '8A5A00'},
    ], cols=3)

    add_heading(doc, '02', '公司与产品概览')
    add_table(doc, [[line] for line in find_lines(lines, ['DIGICUTO', 'Lup', '注册', '产品', '临床', 'AI', '机器人'], 12)], header=['关键信息'])

    add_heading(doc, '03', '监管、商业化与竞争')
    add_table(doc, [[line] for line in find_lines(lines, ['FDA', 'MDR', 'CDSCO', '商业', 'Smile-Link', 'Yomi', 'Perceptive'], 14)], header=['要点'])

    add_heading(doc, '04', '财务、SWOT 与风险矩阵')
    risks = find_lines(lines, ['财务风险', '监管风险', '竞争风险', '商业化风险', '融资', '注册资本'], 10)
    add_table(doc, [[idx + 1, value, '需人工复核'] for idx, value in enumerate(risks)], header=['序号', '风险/财务要点', '状态'])

    add_heading(doc, '05', '来源与下一步')
    add_table(doc, [[line] for line in find_lines(lines, ['来源', '官方', '权威媒体', '学术研究', '免责声明'], 12)], header=['来源/说明'])
    add_callout(doc, '人工复核建议', '请重点复核监管状态、融资与估值、授权协议、商业化数据、供应链稳定性。', fill='EEF2F7')

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Office Agent 专业报告版 | Confidential').font.size = Pt(8)
    doc.save(output)
    return output
