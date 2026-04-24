"""Shared DOCX helpers for generic Office Agent templates."""

import os

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = '1F2933'
SLATE = '52606D'
BLUE = '1D4ED8'
TEAL = '0F766E'
AMBER = 'B45309'
RED = 'B42318'
PAPER = 'F8FAFC'
BLUE_TINT = 'EAF2FF'
TEAL_TINT = 'E6F7F4'
AMBER_TINT = 'FFF4D8'
RED_TINT = 'FCE8E6'
GRAY_TINT = 'F1F5F9'


def set_font(run, size=None, bold=False, color=None, font='Microsoft YaHei'):
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    for name in ['Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3']:
        style = doc.styles[name]
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = RGBColor.from_string(INK)
    doc.styles['Title'].font.size = Pt(24)
    doc.styles['Title'].font.bold = True
    doc.styles['Title'].font.color.rgb = RGBColor.from_string(INK)
    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 1'].font.bold = True
    doc.styles['Heading 1'].font.color.rgb = RGBColor.from_string(BLUE)
    doc.styles['Heading 2'].font.size = Pt(12)
    doc.styles['Heading 2'].font.bold = True
    doc.styles['Heading 2'].font.color.rgb = RGBColor.from_string(TEAL)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)


def add_para(doc, text='', size=None, bold=False, color=None, align=None, after=5, before=0):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, color=None, size=9.3, fill=None):
    if fill:
        shade_cell(cell, fill)
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    set_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, rows, header=None, first_col_fill=BLUE_TINT):
    if not rows and not header:
        return None
    col_count = len(header) if header else max(len(row) for row in rows)
    row_count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = 'Table Grid'
    row_offset = 0
    if header:
        row_offset = 1
        for idx, value in enumerate(header):
            set_cell_text(table.rows[0].cells[idx], value, bold=True, color='FFFFFF', fill=BLUE)
    for r_idx, row in enumerate(rows, start=row_offset):
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else ''
            fill = first_col_fill if c_idx == 0 and first_col_fill else None
            set_cell_text(table.rows[r_idx].cells[c_idx], value, bold=(c_idx == 0), fill=fill, size=9)
    return table


def add_callout(doc, title, body, fill=GRAY_TINT, color=INK):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_text(cell, title, bold=True, color=color, size=10)
    paragraph = cell.add_paragraph(str(body))
    for run in paragraph.runs:
        set_font(run, size=9.2, color=INK)
    return table


def add_badges(doc, badges):
    if not badges:
        return None
    table = doc.add_table(rows=1, cols=len(badges))
    for idx, badge in enumerate(badges):
        set_cell_text(
            table.rows[0].cells[idx],
            badge.get('label', ''),
            bold=True,
            color=badge.get('color', INK),
            fill=badge.get('fill', GRAY_TINT),
            size=9.2,
        )
    return table


def add_cover(doc, title, subtitle, badges=None):
    add_para(doc, 'OFFICE AGENT CONVERSION', size=9.5, bold=True, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_para(doc, title, size=24, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, subtitle, size=10, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    add_badges(doc, badges or [])
    add_para(doc, '', after=8)


def iter_text_items(extracted):
    for slide in extracted.get('slides', []):
        for item in slide.get('items', []):
            if item.get('type') == 'text' and item.get('text'):
                yield slide, item


def slide_texts(slide):
    return [item.get('text', '') for item in slide.get('items', []) if item.get('type') == 'text' and item.get('text')]


def title_from_source(source):
    return os.path.splitext(os.path.basename(source or 'Converted presentation'))[0]


def add_item_image(doc, item, include_images):
    if not include_images or not item.get('path'):
        return False
    try:
        doc.add_picture(item['path'], width=Inches(5.9))
        return True
    except Exception as exc:
        add_para(doc, f'[Image omitted: {exc}]', color=RED, size=9)
        return False
