"""Generic reading-oriented DOCX template."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from .template_common import (
        AMBER_TINT,
        BLUE_TINT,
        GRAY_TINT,
        RED,
        RED_TINT,
        SLATE,
        add_callout,
        add_cover,
        add_item_image,
        add_para,
        add_table,
        slide_texts,
        style_document,
        title_from_source,
    )
except ImportError:
    from template_common import (
        AMBER_TINT,
        BLUE_TINT,
        GRAY_TINT,
        RED,
        RED_TINT,
        SLATE,
        add_callout,
        add_cover,
        add_item_image,
        add_para,
        add_table,
        slide_texts,
        style_document,
        title_from_source,
    )


def collect_summary_rows(extracted, limit=8):
    rows = []
    for slide in extracted.get('slides', []):
        texts = slide_texts(slide)
        if not texts:
            continue
        rows.append([slide.get('slide_no'), slide.get('title') or texts[0][:48], texts[0][:180]])
        if len(rows) >= limit:
            break
    return rows


def semantic_fill(item):
    semantic = item.get('semantic_guess')
    if semantic == 'risk':
        return RED_TINT
    if semantic == 'source_or_disclaimer':
        return GRAY_TINT
    if semantic == 'summary_or_recommendation':
        return AMBER_TINT
    return BLUE_TINT


def render(extracted, plan, output):
    doc = Document()
    style_document(doc)
    source = extracted.get('source', '')
    title = title_from_source(source)

    add_cover(doc, title, '通用阅读版 | 保留原始顺序，提升 Word 可读性', [
        {'label': f"模式: {plan.get('selected_mode', 'generic_reading')}"},
        {'label': f"忠实度: {plan.get('fidelity_level', 'F1/F2')}", 'fill': AMBER_TINT},
        {'label': f"风险级别: {plan.get('risk_level', 'unknown')}", 'fill': RED_TINT, 'color': RED},
    ])
    add_callout(
        doc,
        '转换边界',
        '本版本不按业务类型重写内容，只把 PPT 信息整理为适合阅读的章节、段落、表格和提示框。',
        fill=GRAY_TINT,
    )

    rows = collect_summary_rows(extracted)
    if rows:
        doc.add_heading('内容导航', level=1)
        add_table(doc, rows, header=['页码', '原页标题', '首要内容'], first_col_fill=GRAY_TINT)

    for slide in extracted.get('slides', []):
        doc.add_heading(f"{slide.get('slide_no')}. {slide.get('title') or 'Untitled slide'}", level=1)
        text_count = 0
        for item in slide.get('items', []):
            item_type = item.get('type')
            if item_type == 'text':
                text = item.get('text', '')
                if item.get('semantic_guess') in {'risk', 'source_or_disclaimer', 'summary_or_recommendation'}:
                    add_callout(doc, item.get('semantic_guess', 'Note').replace('_', ' ').title(), text, fill=semantic_fill(item))
                elif item.get('level', 0) > 0:
                    doc.add_paragraph(text, style='List Bullet')
                elif len(text) <= 36 and text_count > 0:
                    doc.add_heading(text, level=2)
                else:
                    add_para(doc, text)
                text_count += 1
            elif item_type == 'table':
                add_table(doc, item.get('rows', []), first_col_fill=BLUE_TINT)
            elif item_type == 'image':
                if add_item_image(doc, item, plan.get('include_images')):
                    add_para(doc, '图片按确认方案嵌入。', size=8.5, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER)
            elif item.get('semantic_guess') == 'complex_visual':
                add_callout(doc, '复杂视觉元素', '该元素已进入 source map / fidelity ledger，建议人工复核其空间关系。', fill=GRAY_TINT)
        if slide.get('notes'):
            doc.add_heading('Speaker Notes', level=2)
            add_para(doc, slide['notes'])

    doc.save(output)
    return output
