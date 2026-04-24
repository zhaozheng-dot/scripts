"""Generic visual-report DOCX template."""

from docx import Document

try:
    from .template_common import (
        AMBER,
        AMBER_TINT,
        BLUE,
        BLUE_TINT,
        GRAY_TINT,
        RED,
        RED_TINT,
        TEAL,
        TEAL_TINT,
        add_callout,
        add_cover,
        add_item_image,
        add_para,
        add_table,
        iter_text_items,
        shade_cell,
        set_cell_text,
        slide_texts,
        style_document,
        title_from_source,
    )
except ImportError:
    from template_common import (
        AMBER,
        AMBER_TINT,
        BLUE,
        BLUE_TINT,
        GRAY_TINT,
        RED,
        RED_TINT,
        TEAL,
        TEAL_TINT,
        add_callout,
        add_cover,
        add_item_image,
        add_para,
        add_table,
        iter_text_items,
        shade_cell,
        set_cell_text,
        slide_texts,
        style_document,
        title_from_source,
    )


def pick_items(extracted, semantic, limit=6):
    picked = []
    for slide, item in iter_text_items(extracted):
        if item.get('semantic_guess') == semantic:
            picked.append([slide.get('slide_no'), item.get('text', '')])
        if len(picked) >= limit:
            break
    return picked


def top_slides(extracted, limit=6):
    scored = []
    for slide in extracted.get('slides', []):
        items = slide.get('items', [])
        score = len(items) + sum(4 for item in items if item.get('semantic_guess') in {'risk', 'summary_or_recommendation'})
        scored.append((score, slide))
    return [slide for _, slide in sorted(scored, key=lambda pair: pair[0], reverse=True)[:limit]]


def add_cards(doc, cards, cols=3):
    if not cards:
        return
    rows = (len(cards) + cols - 1) // cols
    table = doc.add_table(rows=rows, cols=cols)
    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx >= len(cards):
                shade_cell(cell, GRAY_TINT)
                continue
            card = cards[idx]
            set_cell_text(cell, card['title'], bold=True, color=card.get('color', BLUE), fill=card.get('fill', BLUE_TINT), size=9.5)
            paragraph = cell.add_paragraph(card.get('body', ''))
            for run in paragraph.runs:
                run.font.size = None
            idx += 1


def make_cards(extracted):
    slides = top_slides(extracted, limit=6)
    cards = []
    palette = [
        (BLUE_TINT, BLUE),
        (TEAL_TINT, TEAL),
        (AMBER_TINT, AMBER),
        (RED_TINT, RED),
        (GRAY_TINT, BLUE),
        (BLUE_TINT, TEAL),
    ]
    for idx, slide in enumerate(slides):
        texts = slide_texts(slide)
        fill, color = palette[idx % len(palette)]
        cards.append({
            'title': f"{slide.get('slide_no')}. {slide.get('title') or 'Key slide'}",
            'body': '\n'.join(texts[:3])[:420] or 'See source map.',
            'fill': fill,
            'color': color,
        })
    return cards


def render(extracted, plan, output):
    doc = Document()
    style_document(doc)
    source = extracted.get('source', '')
    title = title_from_source(source)

    add_cover(doc, title, '通用视觉报告版 | 将卡片、矩阵、流程转译为 Word 原生组件', [
        {'label': f"模式: {plan.get('selected_mode', 'generic_visual_report')}"},
        {'label': f"忠实度: {plan.get('fidelity_level', 'F2')}", 'fill': AMBER_TINT, 'color': AMBER},
        {'label': '专业模板: 未绑定', 'fill': TEAL_TINT, 'color': TEAL},
    ])
    add_callout(
        doc,
        '转换边界',
        '本版本适合通用展示型 PPT：保留来源事实，重组为摘要卡片、重点矩阵、逐页附录；不引入特定业务判断。',
        fill=GRAY_TINT,
    )

    doc.add_heading('一页摘要', level=1)
    add_cards(doc, make_cards(extracted), cols=3)

    risks = pick_items(extracted, 'risk', limit=8)
    if risks:
        doc.add_heading('风险 / 注意事项', level=1)
        add_table(doc, risks, header=['来源页', '内容'], first_col_fill=RED_TINT)

    recommendations = pick_items(extracted, 'summary_or_recommendation', limit=8)
    if recommendations:
        doc.add_heading('摘要 / 建议', level=1)
        add_table(doc, recommendations, header=['来源页', '内容'], first_col_fill=AMBER_TINT)

    sources = pick_items(extracted, 'source_or_disclaimer', limit=8)
    if sources:
        doc.add_heading('来源 / 免责声明', level=1)
        add_table(doc, sources, header=['来源页', '内容'], first_col_fill=GRAY_TINT)

    doc.add_heading('逐页内容附录', level=1)
    for slide in extracted.get('slides', []):
        doc.add_heading(f"{slide.get('slide_no')}. {slide.get('title') or 'Untitled slide'}", level=2)
        texts = slide_texts(slide)
        if texts:
            add_table(doc, [[text] for text in texts[:12]], header=['文本内容'], first_col_fill=None)
        for item in slide.get('items', []):
            if item.get('type') == 'table':
                add_table(doc, item.get('rows', []), first_col_fill=BLUE_TINT)
            elif item.get('type') == 'image':
                add_item_image(doc, item, plan.get('include_images'))
            elif item.get('semantic_guess') == 'complex_visual':
                add_callout(doc, '复杂视觉元素', '已记录到 fidelity ledger，建议人工复核。', fill=GRAY_TINT)

    doc.save(output)
    return output
